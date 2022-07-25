from ast import If
from cgitb import small
from html.entities import entitydefs
from itertools import count
from msilib.schema import ComboBox, Font
from operator import truediv
from pickle import TRUE
from sre_parse import State
from xml.dom.minidom import Entity
from isort import place_module
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from bs4 import BeautifulSoup
from pykeyboard import PyKeyboard
from pymouse import PyMouse
import re
import requests
import hashlib
import os
import time
import glob
import tkinter as tk
from tkinter import VERTICAL, ttk
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
import smtplib


def setupWin():
    win.title("八字命盤")
    win.option_add("*Font", "微軟正黑體 20")
    win.resizable(0, 0)
    win.iconbitmap(
        "bagua.ico")


def crawler():
    print(year)
    print(isLeapMonth)
    print(isLunar)
    print(isMan)
    global bigFate
    global bigFateText
    global loveGod
    global body
    driver = webdriver.Edge(executable_path="msedgedriver")
    driver.set_window_size(1552, 893)
    driver.get("https://myfate.herokuapp.com/")
    print(isLunar)
    if isLunar:
        driver.find_element("id", "lunar").click()
    if isMan:
        driver.find_element("id", "male").click()
    select = Select(driver.find_element("id", "Year"))
    select.select_by_index(year - 1900)
    select = Select(driver.find_element("id", "Month"))
    select.select_by_index(mon - 1)
    select = Select(driver.find_element("id", "Day"))
    select.select_by_index(day - 1)
    select = Select(driver.find_element("id", "Hour"))
    select.select_by_index(hr)
    if(isLunar and isLeapMonth):
        driver.find_element(By.NAME, "Leap").click()
    driver.find_element(By.TAG_NAME, "button").click()
    element = WebDriverWait(driver, 10).until(
        EC.presence_of_element_located(
            (By.CLASS_NAME, "card-header"))
    )
    print("OK")
    firstData_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'card my-3 shadow'][last()]//div[@class = 'flexcontainer']//span[@class = 'v']")
    for i in range(0, len(firstData_obj)):
        """ print(i)
        print(str((firstData_obj[i].text))) """
        if i > 0 and i < 5:
            mainStar.append(str((firstData_obj[i].text)))
        if i > 5 and i < 18:
            secondStar[(i - 6) // 3].append(str((firstData_obj[i].text)))
        if i > 18 and i < 23:
            fate.append(str((firstData_obj[i].text)))
    secondData_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'card my-3 shadow'][last()]//div[@class = 'flexcontainer']//span[@class = 'vs']")
    for i in range(0, len(secondData_obj)):
        if i > 0:
            bury[(i - 1) // 3].append(str((secondData_obj[i].text)))
    getLen_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'card my-3 shadow'][last()]//div[@class = 'flexcontainer']//div[@class = 'subcontainer'][last()]/div[@class='s']")
    getlen = []
    for i in getLen_obj:
        getlen.append(len(str(i.text).replace('\n', '')))
    third_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'card my-3 shadow'][last()]//div[@class = 'flexcontainer']//div[@class = 'subcontainer'][last()]//span[@class='vl']")
    index = 0
    sumLen = [0, 0, 0, 0]
    allshansha = []
    for i in range(0, len(third_obj)):
        if sumLen[index] == getlen[index]:
            index += 1
        allshansha.append(str((third_obj[i].text)))
        shansha[index].append(str((third_obj[i].text)))
        sumLen[index] += len((str((third_obj[i].text))))
    eightword_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'card my-3 shadow'][last()]//div[@class = 'flexcontainer']//div[@class = 'pillar']")
    for i in range(0, len(eightword_obj)):
        eightword_f.append(str((eightword_obj[i].text))[0])
        eightword_s.append(str((eightword_obj[i].text))[1])
    othershansha_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'card my-3 shadow'][last()]/div[@class = 'pane']/div[@class = 'mx-3']//td")
    for i in range(0, len(othershansha_obj)):
        if i % 2 == 0 and str((othershansha_obj[i].text)) not in allshansha:
            othershansha.append(str((othershansha_obj[i].text)))
    bigFate_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'card my-3 shadow'][last()-1]//div[@class = 'pane']/div[@class = 'flexcontainer'][1]/div[@class = 'subcontainer'][last()-1]/div[@class = 's'][1]/span[1]")
    bigFate = str(bigFate_obj[0].text)
    bigFateText_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'card my-3 shadow'][last()-1]//div[@class = 'pane']/div[@class = 'flexcontainer'][1]/div[@class = 'subcontainer'][last()]")
    bigFateText = str(bigFateText_obj[0].text)
    bigFateYears_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'card my-3 shadow'][last()-1]//div[@class = 'pane']/div[@class = 'flexcontainer'][1]/div[@class = 'subcontainer'][last()-1]/div[@class = 's']/span[@class = 'v']")
    for i in range(0, len(bigFateYears_obj)):
        if i < 10:
            bigFateYears.append(str(bigFateYears_obj[i].text))
    loveGod_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'container px-4'][3]/div[@class = 'card my-3 shadow'][2]/div[@class = 'card-body']")
    loveGod = str(loveGod_obj[0].text).split('\n')[3]
    body_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'container px-4'][3]/div[@class = 'card my-3 shadow'][5]//div[@class = 'pane']/div[@class = 'flexcontainer']/div[@class = 'hcontainer'][3]")
    body = str(body_obj[0].text)
    personality_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'container px-4'][last()-2]//div[@class = 'card my-3 shadow'][last()]//td")
    for i in range(0, len(personality_obj), 2):
        personality[str(personality_obj[i].text)] = str(
            personality_obj[i+1].text)
    horseFlower_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'container px-4'][last()-1]//table[@class ='table table-striped table-bordered']//td")
    for i in range(0, len(horseFlower_obj), 2):
        horseFlower[str(horseFlower_obj[i].text)] = str(
            horseFlower_obj[i+1].text)
    shanshaExplain_obj = driver.find_elements(
        By.XPATH, "//div[@class = 'container px-4'][last()]//table[@class ='table table-striped table-bordered']//td")
    for i in range(0, len(shanshaExplain_obj), 2):
        shanshaExplain[str(shanshaExplain_obj[i].text)] = str(
            shanshaExplain_obj[i+1].text)
    driver.quit()


def finalcheck(l):
    global isLeapMonth
    isLeapMonth = (l == 1)
    win.destroy()


def GUI3():
    global win
    global isLeapMonth
    global isLeapMonth_c
    win = tk.Tk()
    setupWin()

    isLeapMonth_l = tk.Label(text="是閏月嗎", font="微軟正黑體 20")
    isLeapMonth_l.grid(column=0, row=0, sticky="W")
    leapMon = tk.IntVar()
    isLeapMonth_l1 = tk.Radiobutton(
        win, text='是', variable=leapMon, value=1, font="微軟正黑體 20")
    isLeapMonth_l2 = tk.Radiobutton(
        win, text='否', variable=leapMon, value=2, font="微軟正黑體 20")
    isLeapMonth_l1.grid(column=1, row=0, sticky="W")
    isLeapMonth_l2.select()
    isLeapMonth_l2.grid(column=2, row=0, sticky="w")

    isLeapMonth_b = tk.Button(text="確定", command=lambda: finalcheck(
        leapMon.get()), font="微軟正黑體 15")
    isLeapMonth_b.grid(column=2, row=1, sticky="W")


def checkdata(l, g, y, m, d, h):
    global name
    global recorded
    global isLunar
    global isMan
    global year
    global mon
    global day
    global hr
    global checkAll
    isMan = (g == 1)
    isLunar = (l == 2)
    year = y
    mon = m
    day = d
    hr = h
    win.destroy()
    if isLunar:
        GUI3()
    checkAll = True


def GUI2():
    global win
    global isLunar_l1
    global isLunar_l2
    global year_c
    global mon_c
    global day_c
    global hr_c
    global radioValue
    win = tk.Tk()
    setupWin()
    radioValue = tk.IntVar()
    isLunar_l1 = tk.Radiobutton(
        win, text='陽歷', variable=radioValue, value=1, font="微軟正黑體 20")
    isLunar_l2 = tk.Radiobutton(
        win, text='農曆', variable=radioValue, value=2, font="微軟正黑體 20")
    isLunar_l1.grid(column=0, row=2, sticky="W")
    isLunar_l1.select()
    isLunar_l2.grid(column=1, row=2, sticky="w")

    gender = tk.IntVar()
    isman_l1 = tk.Radiobutton(
        win, text='男', variable=gender, value=1, font="微軟正黑體 20")
    isman_l2 = tk.Radiobutton(
        win, text='女', variable=gender, value=2, font="微軟正黑體 20")
    isman_l1.grid(column=0, row=3, sticky="W")
    isman_l1.select()
    isman_l2.grid(column=1, row=3, sticky="w")

    year_l1 = tk.Label(text="西元", font="微軟正黑體 20")
    year_l2 = tk.Label(text="年", font="微軟正黑體 20")
    year_c = ttk.Combobox(win, state="readonly", font="微軟正黑體 20")
    years = []
    for i in range(1900, 2051):
        years.append(f" {i}")
    year_c['value'] = years
    year_c.current(100)
    year_l1.grid(column=0, row=4, sticky="W")
    year_c.grid(column=1, row=4, sticky="W")
    year_l2.grid(column=2, row=4, sticky="W")

    mon_l = tk.Label(text="月", font="微軟正黑體 20")
    mon_c = ttk.Combobox(win, state="readonly", font="微軟正黑體 20")
    mons = []
    for i in range(1, 13):
        mons.append(f" {i}")
    mon_c['value'] = mons
    mon_c.current(0)
    mon_c.grid(column=1, row=5, sticky="W")
    mon_l.grid(column=2, row=5, sticky="W")

    day_l = tk.Label(text="日", font="微軟正黑體 20")
    day_c = ttk.Combobox(win, state="readonly", font="微軟正黑體 20")
    days = []
    for i in range(1, 32):
        days.append(f" {i}")
    day_c['value'] = days
    day_c.current(0)
    day_c.grid(column=1, row=6, sticky="W")
    day_l.grid(column=2, row=6, sticky="W")

    hr_l = tk.Label(text="時", font="微軟正黑體 20")
    hr_c = ttk.Combobox(win, state="readonly", font="微軟正黑體 20")
    hrs = []
    for i in range(0, 24):
        hrs.append(f" {i}")
    hr_c['value'] = hrs
    hr_c.current(0)
    hr_c.grid(column=1, row=7, sticky="W")
    hr_l.grid(column=2, row=7, sticky="W")

    data_b = tk.Button(text="確定", command=lambda: checkdata(radioValue.get(), gender.get(), int(
        year_c.get()), int(mon_c.get()), int(day_c.get()), int(hr_c.get())), font="微軟正黑體 15")
    data_b.grid(row=8, column=3)

    win.mainloop()


def checkname(n):
    global recData
    global name
    global recorded
    global isLunar
    global isMan
    global year
    global mon
    global day
    global hr
    global isLeapMonth
    global checkAll
    name = n

    with open("names.txt", 'r', encoding="utf-8") as f:
        names = f.readlines()
        for i in names:
            recData.append(i.split(' ')[:-1])
            if name == i.split(' ')[1]:
                recorded = TRUE
                year = int(i.split(' ')[0])
                isLunar = (i.split(' ')[2] == "True")
                isMan = (i.split(' ')[3] == "True")
                mon = int(i.split(' ')[4])
                day = int(i.split(' ')[5])
                hr = int(i.split(' ')[6])
                isLeapMonth = bool(i.split(' ')[7] == "True")

    win.destroy()
    if recorded:
        checkAll = True
        return
    GUI2()


def GUI():
    global name
    global visited
    global win
    win.destroy()
    win = tk.Tk()
    setupWin()
    name_l = tk.Label(text="姓名", font="微軟正黑體 20")
    name_c = ttk.Combobox(win, font="微軟正黑體 20")
    recnames = []
    with open("names.txt", 'r', encoding="utf-8") as f:
        names = f.readlines()
        for i in names:
            recnames.append(i.split(' ')[1])
    name_c['value'] = recnames
    name_c.current(0)
    name_b = tk.Button(
        text="確定", command=lambda: checkname(name_c.get()), font="微軟正黑體 15")
    name_l.grid(row=0, column=0)
    name_c.grid(row=0, column=1)
    name_b.grid(row=0, column=2)

    win.mainloop()


def addNewData():
    global recData
    global name
    global recorded
    global isLunar
    global isMan
    global year
    global mon
    global day
    global hr
    global isLeapMonth
    newData = []
    if not recorded and checkAll:
        newData = [year, name, isLunar, isMan, mon, day, hr, isLeapMonth]
        recData.append(newData)
        recData.sort(key=lambda x: int(x[0]))
        with open("names.txt", 'w', encoding="utf-8") as f:
            for Data in recData:
                for i in Data:
                    f.write(f"{i} ")
                f.write('\n')


def five_g(i):
    if(i == 1 or i == 0):
        return "木"
    if(i == 3 or i == 2):
        return "火"
    if(i == 5 or i == 4):
        return "土"
    if(i == 7 or i == 6):
        return "金"
    if(i == 9 or i == 8):
        return "水"


def five_z(i):
    if(i == 3 or i == 2):
        return "木"
    if(i == 6 or i == 5):
        return "火"
    if(i == 4 or i == 10 or i == 1 or i == 7):
        return "土"
    if(i == 8 or i == 9):
        return "金"
    if(i == 0 or i == 11):
        return "水"


def adjust():
    a = ['甲', '乙', '丙', '丁', '戊', '己', '庚', '辛', '壬', '癸']
    b = ['子', '丑', '寅', '卯', '辰', '巳', '午', '未', '申', '酉', '戌', '亥']
    e = []
    for i in range(60):
        e.append(a[i % 10]+b[i % 12])
    tmp = eightword_f[0]+eightword_s[0]
    for i in range(60):
        if tmp == e[i]:
            for j in range(100):
                flowYears.append(e[(i+j) % 60])
            break
    c = {}
    d = {}
    for i in range(10):
        c[a[i]] = i
    for i in range(12):
        d[b[i]] = i
    g = []
    z = []
    for i in range(4):
        g.append(c[eightword_f[i]])
        z.append(d[eightword_s[i]])
        eightword_f_elem.append(five_g(g[i]))
        eightword_s_elem.append(five_z(z[i]))
    listy = []
    listm = []
    listd = []
    listh = []
    if z[1] == 0:
        if z[2] == 5:
            listd.append("天德")
        elif g[2] == 8:
            listd.append("月德")
    if z[1] == 1:
        if g[2] == 6:
            listd.append("天德")
        elif g[2] == 6:
            listd.append("月德")
    if z[1] == 2:
        if g[2] == 3:
            listd.append("天德")
        elif g[2] == 2:
            listd.append("月德")
    if z[1] == 3:
        if z[2] == 8:
            listd.append("天德")
        elif g[2] == 0:
            listd.append("月德")
    if z[1] == 4:
        if g[2] == 8:
            listd.append("天德")
        elif g[2] == 8:
            listd.append("月德")
    if z[1] == 5:
        if g[2] == 7:
            listd.append("天德")
        elif g[2] == 6:
            listd.append("月德")
    if z[1] == 6:
        if z[2] == 11:
            listd.append("天德")
        elif g[2] == 2:
            listd.append("月德")
    if z[1] == 7:
        if g[2] == 0:
            listd.append("天德")
        elif g[2] == 0:
            listd.append("月德")
    if z[1] == 8:
        if g[2] == 9:
            listd.append("天德")
        elif g[2] == 8:
            listd.append("月德")
    if z[1] == 9:
        if z[2] == 2:
            listd.append("天德")
        elif g[2] == 6:
            listd.append("月德")
    if z[1] == 10:
        if g[2] == 2:
            listd.append("天德")
        elif g[2] == 2:
            listd.append("月德")
    if z[1] == 11:
        if g[2] == 1:
            listd.append("天德")
        elif g[2] == 0:
            listd.append("月德")

    if z[0] == 0 or z[0] == 1 or z[0] == 11:
        for i in range(4):
            if z[i] == 2:
                if i == 0:
                    listy.append("孤辰")
                elif i == 1:
                    listm.append("孤辰")
                elif i == 2:
                    listd.append("孤辰")
                else:
                    listh.append("孤辰")
            if z[i] == 10:
                if i == 0:
                    listy.append("寡宿")
                elif i == 1:
                    listm.append("寡宿")
                elif i == 2:
                    listd.append("寡宿")
                else:
                    listh.append("寡宿")
    if z[0] == 2 or z[0] == 3 or z[0] == 4:
        for i in range(4):
            if z[i] == 5:
                if i == 0:
                    listy.append("孤辰")
                elif i == 1:
                    listm.append("孤辰")
                elif i == 2:
                    listd.append("孤辰")
                else:
                    listh.append("孤辰")
            if z[i] == 1:
                if i == 0:
                    listy.append("寡宿")
                elif i == 1:
                    listm.append("寡宿")
                elif i == 2:
                    listd.append("寡宿")
                else:
                    listh.append("寡宿")
    if z[0] == 5 or z[0] == 6 or z[0] == 7:
        for i in range(4):
            if z[i] == 8:
                if i == 0:
                    listy.append("孤辰")
                elif i == 1:
                    listm.append("孤辰")
                elif i == 2:
                    listd.append("孤辰")
                else:
                    listh.append("孤辰")
            if z[i] == 4:
                if i == 0:
                    listy.append("寡宿")
                elif i == 1:
                    listm.append("寡宿")
                elif i == 2:
                    listd.append("寡宿")
                else:
                    listh.append("寡宿")
    if z[0] == 8 or z[0] == 9 or z[0] == 10:
        for i in range(4):
            if z[i] == 11:
                if i == 0:
                    listy.append("孤辰")
                elif i == 1:
                    listm.append("孤辰")
                elif i == 2:
                    listd.append("孤辰")
                else:
                    listh.append("孤辰")
            if z[i] == 7:
                if i == 0:
                    listy.append("寡宿")
                elif i == 1:
                    listm.append("寡宿")
                elif i == 2:
                    listd.append("寡宿")
                else:
                    listh.append("寡宿")

    if z[0] == 0 or z[0] == 4 or z[0] == 8:
        if z[1] == 2:
            listm.append("男掃女家")

        if z[1] == 1:
            listm.append("女掃男家")
    if z[0] == 1 or z[0] == 5 or z[0] == 9:
        if z[1] == 7:
            listm.append("男掃女家")
        if z[1] == 10:
            listm.append("女掃男家")
    if z[0] == 2 or z[0] == 6 or z[0] == 10:
        if z[1] == 5:
            listm.append("男掃女家")
        if z[1] == 8:
            listm.append("女掃男家")
    if z[0] == 3 or z[0] == 7 or z[0] == 11:
        if z[1] == 3:
            listm.append("男掃女家")
        if z[1] == 9:
            listm.append("女掃男家")

    if 0 in g:
        if 4 in g:
            if 6 in g:
                listy.append("天上三奇")
    if 1 in g:
        if 2 in g:
            if 3 in g:
                listy.append("地上三奇")
    if 7 in g:
        if 8 in g:
            if 9 in g:
                listy.append("人中三奇")

    if g[2] == 6 and z[2] == 4:
        listy.append("魁罡")
    if g[2] == 8 and z[2] == 4:
        listy.append("魁罡")
    if g[2] == 6 and z[2] == 10:
        listy.append("魁罡")
    if g[2] == 4 and z[2] == 10:
        listy.append("魁罡")

    if g[2] == 1 and z[2] == 5:
        listy.append("孤鸞日")
    if g[2] == 3 and z[2] == 5:
        listy.append("孤鸞日")
    if g[2] == 0 and z[2] == 2:
        listy.append("孤鸞日")
    if g[2] == 8 and z[2] == 2:
        listy.append("孤鸞日")
    if g[2] == 2 and z[2] == 6:
        listy.append("孤鸞日")
    if g[2] == 4 and z[2] == 6:
        listy.append("孤鸞日")
    if g[2] == 8 and z[2] == 0:
        listy.append("孤鸞日")
    if g[2] == 7 and z[2] == 11:
        listy.append("孤鸞日")
    if g[2] == 4 and z[2] == 8:
        listy.append("孤鸞日")
    shansha[0] = shansha[0] + listy
    shansha[1] = shansha[1] + listm
    shansha[2] = shansha[2] + listd
    shansha[3] = shansha[3] + listh
    listy.clear()
    listm.clear()
    listd.clear()
    listh.clear()
    # 小兒關煞

    for i in range(4):
        if z[i] == 5:
            if g[2] == 0 or g[2] == 5:
                if i == 0:
                    listy.append("落井關")
                elif i == 1:
                    listm.append("落井關")
                elif i == 2:
                    listd.append("落井關")
                else:
                    listh.append("落井關")
        if z[i] == 0:
            if g[2] == 1 or g[2] == 6:
                if i == 0:
                    listy.append("落井關")
                elif i == 1:
                    listm.append("落井關")
                elif i == 2:
                    listd.append("落井關")
                else:
                    listh.append("落井關")
        if z[i] == 8:
            if g[2] == 2 or g[2] == 7:
                if i == 0:
                    listy.append("落井關")
                elif i == 1:
                    listm.append("落井關")
                elif i == 2:
                    listd.append("落井關")
                else:
                    listh.append("落井關")
        if z[i] == 10:
            if g[2] == 3 or g[2] == 8:
                if i == 0:
                    listy.append("落井關")
                elif i == 1:
                    listm.append("落井關")
                elif i == 2:
                    listd.append("落井關")
                else:
                    listh.append("落井關")
        if z[i] == 3:
            if g[2] == 4 or g[2] == 9:
                if i == 0:
                    listy.append("落井關")
                elif i == 1:
                    listm.append("落井關")
                elif i == 2:
                    listd.append("落井關")
                else:
                    listh.append("落井關")

    for i in range(4):
        if z[i] == 1 or z[i] == 9 or z[i] == 5:
            if g[2] == 0 or g[2] == 5:
                if i == 0:
                    listy.append("雞飛關")
                elif i == 1:
                    listm.append("雞飛關")
                elif i == 2:
                    listd.append("雞飛關")
                else:
                    listh.append("雞飛關")
        if z[i] == 0:
            if g[2] == 1 or g[2] == 2 or g[2] == 3 or g[2] == 4:
                if i == 0:
                    listy.append("雞飛關")
                elif i == 1:
                    listm.append("雞飛關")
                elif i == 2:
                    listd.append("雞飛關")
                else:
                    listh.append("雞飛關")
        if z[i] == 7 or z[i] == 9 or z[i] == 11:
            if g[2] == 6:
                if i == 0:
                    listy.append("雞飛關")
                elif i == 1:
                    listm.append("雞飛關")
                elif i == 2:
                    listd.append("雞飛關")
                else:
                    listh.append("雞飛關")
        if z[i] == 10 or z[i] == 6 or z[i] == 2:
            if g[2] == 7 or g[2] == 8 or g[2] == 9:
                if i == 0:
                    listy.append("雞飛關")
                elif i == 1:
                    listm.append("雞飛關")
                elif i == 2:
                    listd.append("雞飛關")
                else:
                    listh.append("雞飛關")

    for i in range(4):
        if z[i] == 4 or z[i] == 0 or z[i] == 8:
            if g[2] == 0 or g[2] == 1 or g[2] == 2 or g[2] == 3:
                if i == 0:
                    listy.append("取命關")
                elif i == 1:
                    listm.append("取命關")
                elif i == 2:
                    listd.append("取命關")
                else:
                    listh.append("取命關")
        if z[i] == 7 or z[i] == 3 or z[i] == 11:
            if g[2] == 4 or g[2] == 5 or g[2] == 6:
                if i == 0:
                    listy.append("取命關")
                elif i == 1:
                    listm.append("取命關")
                elif i == 2:
                    listd.append("取命關")
                else:
                    listh.append("取命關")
        if z[i] == 10 or z[i] == 6 or z[i] == 2:
            if g[2] == 7 or g[2] == 8 or g[2] == 9:
                if i == 0:
                    listy.append("取命關")
                elif i == 1:
                    listm.append("取命關")
                elif i == 2:
                    listd.append("取命關")
                else:
                    listh.append("取命關")

    for i in range(4):
        if z[i] == 1 or z[i] == 6:
            if g[2] == 0 or g[2] == 1:
                if i == 0:
                    listy.append("雷公關")
                elif i == 1:
                    listm.append("雷公關")
                elif i == 2:
                    listd.append("雷公關")
                else:
                    listh.append("雷公關")
        if z[i] == 0:
            if g[2] == 2 or g[2] == 3 or g[2] == 8 or g[2] == 9:
                if i == 0:
                    listy.append("雷公關")
                elif i == 1:
                    listm.append("雷公關")
                elif i == 2:
                    listd.append("雷公關")
                else:
                    listh.append("雷公關")
        if z[i] == 10:
            if g[2] == 4:
                if i == 0:
                    listy.append("雷公關")
                elif i == 1:
                    listm.append("雷公關")
                elif i == 2:
                    listd.append("雷公關")
                else:
                    listh.append("雷公關")
        if z[i] == 6:
            if g[2] == 5:
                if i == 0:
                    listy.append("雷公關")
                elif i == 1:
                    listm.append("雷公關")
                elif i == 2:
                    listd.append("雷公關")
                else:
                    listh.append("雷公關")
        if z[i] == 2:
            if g[2] == 6 or g[2] == 7:
                if i == 0:
                    listy.append("雷公關")
                elif i == 1:
                    listm.append("雷公關")
                elif i == 2:
                    listd.append("雷公關")
                else:
                    listh.append("雷公關")

    for i in range(4):
        if z[i] == 7 or z[i] == 8:
            if g[2] == 0 or g[2] == 1:
                if i == 0:
                    listy.append("斷腸關")
                elif i == 1:
                    listm.append("斷腸關")
                elif i == 2:
                    listd.append("斷腸關")
                else:
                    listh.append("斷腸關")
        if z[i] == 4 or z[i] == 5:
            if g[2] == 2 or g[2] == 3:
                if i == 0:
                    listy.append("斷腸關")
                elif i == 1:
                    listm.append("斷腸關")
                elif i == 2:
                    listd.append("斷腸關")
                else:
                    listh.append("斷腸關")
        if z[i] == 2:
            if g[2] == 6 or g[2] == 7:
                if i == 0:
                    listy.append("斷腸關")
                elif i == 1:
                    listm.append("斷腸關")
                elif i == 2:
                    listd.append("斷腸關")
                else:
                    listh.append("斷腸關")
        if z[i] == 1:
            if g[2] == 8 or g[2] == 9:
                if i == 0:
                    listy.append("斷腸關")
                elif i == 1:
                    listm.append("斷腸關")
                elif i == 2:
                    listd.append("斷腸關")
                else:
                    listh.append("斷腸關")

    for i in range(4):
        if z[i] == 6 or z[i] == 7:
            if g[2] == 0 or g[2] == 1:
                if i == 0:
                    listy.append("千日關")
                elif i == 1:
                    listm.append("千日關")
                elif i == 2:
                    listd.append("千日關")
                else:
                    listh.append("千日關")
        if z[i] == 8 or z[i] == 9:
            if g[2] == 2 or g[2] == 3:
                if i == 0:
                    listy.append("千日關")
                elif i == 1:
                    listm.append("千日關")
                elif i == 2:
                    listd.append("千日關")
                else:
                    listh.append("千日關")
        if z[i] == 10 or z[i] == 5:
            if g[2] == 4 or g[2] == 5:
                if i == 0:
                    listy.append("千日關")
                elif i == 1:
                    listm.append("千日關")
                elif i == 2:
                    listd.append("千日關")
                else:
                    listh.append("千日關")
        if z[i] == 2:
            if g[2] == 6 or g[2] == 7:
                if i == 0:
                    listy.append("千日關")
                elif i == 1:
                    listm.append("千日關")
                elif i == 2:
                    listd.append("千日關")
                else:
                    listh.append("千日關")
        if z[i] == 11 or z[i] == 1:
            if g[2] == 8 or g[2] == 9:
                if i == 0:
                    listy.append("千日關")
                elif i == 1:
                    listm.append("千日關")
                elif i == 2:
                    listd.append("千日關")
                else:
                    listh.append("千日關")

    for i in range(4):
        if z[i] == 8 or z[i] == 9:
            if g[2] == 0 or g[2] == 1:
                if i == 0:
                    listy.append("急腳關")
                elif i == 1:
                    listm.append("急腳關")
                elif i == 2:
                    listd.append("急腳關")
                else:
                    listh.append("急腳關")
        if z[i] == 0 or z[i] == 11:
            if g[2] == 2 or g[2] == 3:
                if i == 0:
                    listy.append("急腳關")
                elif i == 1:
                    listm.append("急腳關")
                elif i == 2:
                    listd.append("急腳關")
                else:
                    listh.append("急腳關")
        if z[i] == 3 or z[i] == 2:
            if g[2] == 4 or g[2] == 5:
                if i == 0:
                    listy.append("急腳關")
                elif i == 1:
                    listm.append("急腳關")
                elif i == 2:
                    listd.append("急腳關")
                else:
                    listh.append("急腳關")
        if z[i] == 6 or z[i] == 5:
            if g[2] == 6 or g[2] == 7:
                if i == 0:
                    listy.append("急腳關")
                elif i == 1:
                    listm.append("急腳關")
                elif i == 2:
                    listd.append("急腳關")
                else:
                    listh.append("急腳關")
        if z[i] == 7 or z[i] == 1 or z[i] == 10 or z[i] == 4:
            if g[2] == 8 or g[2] == 9:
                if i == 0:
                    listy.append("急腳關")
                elif i == 1:
                    listm.append("急腳關")
                elif i == 2:
                    listd.append("急腳關")
                else:
                    listh.append("急腳關")

    for i in range(4):
        if z[i] == 8 or z[i] == 7:
            if g[2] == 2 or g[2] == 3:
                if i == 0:
                    listy.append("鐵蛇關")
                elif i == 1:
                    listm.append("鐵蛇關")
                elif i == 2:
                    listd.append("鐵蛇關")
                else:
                    listh.append("鐵蛇關")
        if z[i] == 4:
            if g[2] == 0 or g[2] == 1:
                if i == 0:
                    listy.append("鐵蛇關")
                elif i == 1:
                    listm.append("鐵蛇關")
                elif i == 2:
                    listd.append("鐵蛇關")
                else:
                    listh.append("鐵蛇關")
        if z[i] == 2:
            if g[2] == 4 or g[2] == 5:
                if i == 0:
                    listy.append("鐵蛇關")
                elif i == 1:
                    listm.append("鐵蛇關")
                elif i == 2:
                    listd.append("鐵蛇關")
                else:
                    listh.append("鐵蛇關")
        if z[i] == 10:
            if g[2] == 6 or g[2] == 7:
                if i == 0:
                    listy.append("鐵蛇關")
                elif i == 1:
                    listm.append("鐵蛇關")
                elif i == 2:
                    listd.append("鐵蛇關")
                else:
                    listh.append("鐵蛇關")
        if z[i] == 1:
            if g[2] == 8 or g[2] == 9:
                if i == 0:
                    listy.append("鐵蛇關")
                elif i == 1:
                    listm.append("鐵蛇關")
                elif i == 2:
                    listd.append("鐵蛇關")
                else:
                    listh.append("鐵蛇關")

    for i in range(4):
        if z[i] == 9:
            if g[2] == 0 or g[2] == 1:
                if i == 0:
                    listy.append("白虎關")
                elif i == 1:
                    listm.append("白虎關")
                elif i == 2:
                    listd.append("白虎關")
                else:
                    listh.append("白虎關")
        if z[i] == 0:
            if g[2] == 2 or g[2] == 3:
                if i == 0:
                    listy.append("白虎關")
                elif i == 1:
                    listm.append("白虎關")
                elif i == 2:
                    listd.append("白虎關")
                else:
                    listh.append("白虎關")
        if z[i] == 6:
            if g[2] == 4 or g[2] == 5 or g[2] == 8 or g[2] == 9:
                if i == 0:
                    listy.append("白虎關")
                elif i == 1:
                    listm.append("白虎關")
                elif i == 2:
                    listd.append("白虎關")
                else:
                    listh.append("白虎關")
        if z[i] == 3:
            if g[2] == 6 or g[2] == 7:
                if i == 0:
                    listy.append("白虎關")
                elif i == 1:
                    listm.append("白虎關")
                elif i == 2:
                    listd.append("白虎關")
                else:
                    listh.append("白虎關")

    shansha[3] = shansha[3] + listh
    listh.clear()

    for i in range(4):
        if z[i] == 9:
            if z[2] == 0:
                if i == 0:
                    listy.append("鬼門關")
                elif i == 1:
                    listm.append("鬼門關")
                elif i == 2:
                    listd.append("鬼門關")
                else:
                    listh.append("鬼門關")
        if z[i] == 6:
            if z[2] == 1:
                if i == 0:
                    listy.append("鬼門關")
                elif i == 1:
                    listm.append("鬼門關")
                elif i == 2:
                    listd.append("鬼門關")
                else:
                    listh.append("鬼門關")
        if z[i] == 7:
            if z[2] == 2:
                if i == 0:
                    listy.append("鬼門關")
                elif i == 1:
                    listm.append("鬼門關")
                elif i == 2:
                    listd.append("鬼門關")
                else:
                    listh.append("鬼門關")
        if z[i] == 8:
            if z[2] == 3:
                if i == 0:
                    listy.append("鬼門關")
                elif i == 1:
                    listm.append("鬼門關")
                elif i == 2:
                    listd.append("鬼門關")
                else:
                    listh.append("鬼門關")
        if z[i] == 11:
            if z[2] == 4:
                if i == 0:
                    listy.append("鬼門關")
                elif i == 1:
                    listm.append("鬼門關")
                elif i == 2:
                    listd.append("鬼門關")
                else:
                    listh.append("鬼門關")
        if z[i] == 10:
            if z[2] == 5:
                if i == 0:
                    listy.append("鬼門關")
                elif i == 1:
                    listm.append("鬼門關")
                elif i == 2:
                    listd.append("鬼門關")
                else:
                    listh.append("鬼門關")
        if z[i] == 1:
            if z[2] == 6:
                if i == 0:
                    listy.append("鬼門關")
                elif i == 1:
                    listm.append("鬼門關")
                elif i == 2:
                    listd.append("鬼門關")
                else:
                    listh.append("鬼門關")
        if z[i] == 2:
            if z[2] == 7:
                if i == 0:
                    listy.append("鬼門關")
                elif i == 1:
                    listm.append("鬼門關")
                elif i == 2:
                    listd.append("鬼門關")
                else:
                    listh.append("鬼門關")
        if z[i] == 3:
            if z[2] == 8:
                if i == 0:
                    listy.append("鬼門關")
                elif i == 1:
                    listm.append("鬼門關")
                elif i == 2:
                    listd.append("鬼門關")
                else:
                    listh.append("鬼門關")
        if z[i] == 0:
            if z[2] == 9:
                if i == 0:
                    listy.append("鬼門關")
                elif i == 1:
                    listm.append("鬼門關")
                elif i == 2:
                    listd.append("鬼門關")
                else:
                    listh.append("鬼門關")
        if z[i] == 5:
            if z[2] == 10:
                if i == 0:
                    listy.append("鬼門關")
                elif i == 1:
                    listm.append("鬼門關")
                elif i == 2:
                    listd.append("鬼門關")
                else:
                    listh.append("鬼門關")
        if z[i] == 4:
            if z[2] == 11:
                if i == 0:
                    listy.append("鬼門關")
                elif i == 1:
                    listm.append("鬼門關")
                elif i == 2:
                    listd.append("鬼門關")
                else:
                    listh.append("鬼門關")

    if z[2] == (z[3]+4) % 12:
        listh.append("五鬼關")
    if z[2] == (z[3]+10) % 12:
        listh.append("天狗煞")

    for i in range(4):
        if z[i] == 1:
            if z[2] == 0 or z[2] == 3 or z[2] == 6 or z[2] == 9:
                if i == 0:
                    listy.append("埋兒關")
                elif i == 1:
                    listm.append("埋兒關")
                elif i == 2:
                    listd.append("埋兒關")
                else:
                    listh.append("埋兒關")
        if z[i] == 3:
            if z[2] == 1 or z[2] == 4 or z[2] == 7 or z[2] == 10:
                if i == 0:
                    listy.append("埋兒關")
                elif i == 1:
                    listm.append("埋兒關")
                elif i == 2:
                    listd.append("埋兒關")
                else:
                    listh.append("埋兒關")
        if z[i] == 8:
            if z[2] == 2 or z[2] == 5 or z[2] == 8 or z[2] == 11:
                if i == 0:
                    listy.append("埋兒關")
                elif i == 1:
                    listm.append("埋兒關")
                elif i == 2:
                    listd.append("埋兒關")
                else:
                    listh.append("埋兒關")

    for i in range(4):
        if z[i] == 5 or z[i] == 6:
            if z[2] == 0 or z[2] == 4 or z[2] == 8:
                if i == 0:
                    listy.append("天吊關")
                elif i == 1:
                    listm.append("天吊關")
                elif i == 2:
                    listd.append("天吊關")
                else:
                    listh.append("天吊關")
        if z[i] == 0 or z[i] == 3:
            if z[2] == 1 or z[2] == 5 or z[2] == 9:
                if i == 0:
                    listy.append("天吊關")
                elif i == 1:
                    listm.append("天吊關")
                elif i == 2:
                    listd.append("天吊關")
                else:
                    listh.append("天吊關")
        if z[i] == 6 or z[i] == 4:
            if z[2] == 2 or z[2] == 6 or z[2] == 10:
                if i == 0:
                    listy.append("天吊關")
                elif i == 1:
                    listm.append("天吊關")
                elif i == 2:
                    listd.append("天吊關")
                else:
                    listh.append("天吊關")
        if z[i] == 8 or z[i] == 6:
            if z[2] == 3 or z[2] == 7 or z[2] == 11:
                if i == 0:
                    listy.append("天吊關")
                elif i == 1:
                    listm.append("天吊關")
                elif i == 2:
                    listd.append("天吊關")
                else:
                    listh.append("天吊關")

    for i in range(4):
        if z[i] == 5:
            if z[2] == 0 or z[2] == 4 or z[2] == 8:
                if i == 0:
                    listy.append("短命關")
                elif i == 1:
                    listm.append("短命關")
                elif i == 2:
                    listd.append("短命關")
                else:
                    listh.append("短命關")
        if z[i] == 2:
            if z[2] == 1 or z[2] == 5 or z[2] == 9:
                if i == 0:
                    listy.append("短命關")
                elif i == 1:
                    listm.append("短命關")
                elif i == 2:
                    listd.append("短命關")
                else:
                    listh.append("短命關")
        if z[i] == 4:
            if z[2] == 2 or z[2] == 6 or z[2] == 10:
                if i == 0:
                    listy.append("短命關")
                elif i == 1:
                    listm.append("短命關")
                elif i == 2:
                    listd.append("短命關")
                else:
                    listh.append("短命關")
        if z[i] == 7:
            if z[2] == 3 or z[2] == 7 or z[2] == 11:
                if i == 0:
                    listy.append("短命關")
                elif i == 1:
                    listm.append("短命關")
                elif i == 2:
                    listd.append("短命關")
                else:
                    listh.append("短命關")

    for i in range(4):
        if z[i] == 6 or z[i] == 1 or z[i] == 10 or z[i] == 4:
            if z[2] == 0 or z[2] == 3 or z[2] == 6 or z[2] == 9:
                if i == 0:
                    listy.append("和尚關")
                elif i == 1:
                    listm.append("和尚關")
                elif i == 2:
                    listd.append("和尚關")
                else:
                    listh.append("和尚關")
        if z[i] == 9 or z[i] == 3 or z[i] == 6 or z[i] == 0:
            if z[2] == 1 or z[2] == 4 or z[2] == 7 or z[2] == 10:
                if i == 0:
                    listy.append("和尚關")
                elif i == 1:
                    listm.append("和尚關")
                elif i == 2:
                    listd.append("和尚關")
                else:
                    listh.append("和尚關")
        if z[i] == 11 or z[i] == 5 or z[i] == 8 or z[i] == 2:
            if z[2] == 2 or z[2] == 5 or z[2] == 8 or z[2] == 11:
                if i == 0:
                    listy.append("和尚關")
                elif i == 1:
                    listm.append("和尚關")
                elif i == 2:
                    listd.append("和尚關")
                else:
                    listh.append("和尚關")

    for i in range(4):
        if z[i] == 6:
            if z[2] == 0 or z[2] == 3 or z[2] == 6 or z[2] == 9:
                if i == 0:
                    listy.append("湯火關")
                elif i == 1:
                    listm.append("湯火關")
                elif i == 2:
                    listd.append("湯火關")
                else:
                    listh.append("湯火關")
        if z[i] == 6:
            if z[2] == 1 or z[2] == 4 or z[2] == 7 or z[2] == 10:
                if i == 0:
                    listy.append("湯火關")
                elif i == 1:
                    listm.append("湯火關")
                elif i == 2:
                    listd.append("湯火關")
                else:
                    listh.append("湯火關")
        if z[i] == 2:
            if z[2] == 2 or z[2] == 5 or z[2] == 8 or z[2] == 11:
                if i == 0:
                    listy.append("湯火關")
                elif i == 1:
                    listm.append("湯火關")
                elif i == 2:
                    listd.append("湯火關")
                else:
                    listh.append("湯火關")

    for i in range(4):
        if z[i] == 2:
            if z[2] == 0 or z[2] == 3 or z[2] == 6 or z[2] == 9:
                if i == 0:
                    listy.append("夜啼關")
                elif i == 1:
                    listm.append("夜啼關")
                elif i == 2:
                    listd.append("夜啼關")
                else:
                    listh.append("夜啼關")
        if z[i] == 7:
            if z[2] == 1 or z[2] == 4 or z[2] == 7 or z[2] == 10:
                if i == 0:
                    listy.append("夜啼關")
                elif i == 1:
                    listm.append("夜啼關")
                elif i == 2:
                    listd.append("夜啼關")
                else:
                    listh.append("夜啼關")
        if z[i] == 9:
            if z[2] == 2 or z[2] == 5 or z[2] == 8 or z[2] == 11:
                if i == 0:
                    listy.append("夜啼關")
                elif i == 1:
                    listm.append("夜啼關")
                elif i == 2:
                    listd.append("夜啼關")
                else:
                    listh.append("夜啼關")

    for i in range(4):
        if z[i] == 5:
            if z[2] == 0 or z[2] == 2:
                if i == 0:
                    listy.append("撞命關")
                elif i == 1:
                    listm.append("撞命關")
                elif i == 2:
                    listd.append("撞命關")
                else:
                    listh.append("撞命關")
        if z[i] == 6:
            if z[2] == 1 or z[2] == 10:
                if i == 0:
                    listy.append("撞命關")
                elif i == 1:
                    listm.append("撞命關")
                elif i == 2:
                    listd.append("撞命關")
                else:
                    listh.append("撞命關")
        if z[i] == 0:
            if z[2] == 3:
                if i == 0:
                    listy.append("撞命關")
                elif i == 1:
                    listm.append("撞命關")
                elif i == 2:
                    listd.append("撞命關")
                else:
                    listh.append("撞命關")
        if z[i] == 6:
            if z[2] == 4 or z[2] == 5 or z[2] == 8:
                if i == 0:
                    listy.append("撞命關")
                elif i == 1:
                    listm.append("撞命關")
                elif i == 2:
                    listd.append("撞命關")
                else:
                    listh.append("撞命關")
        if z[i] == 7:
            if z[2] == 6 or z[2] == 7:
                if i == 0:
                    listy.append("撞命關")
                elif i == 1:
                    listm.append("撞命關")
                elif i == 2:
                    listd.append("撞命關")
                else:
                    listh.append("撞命關")
        if z[i] == 11:
            if z[2] == 9 or z[2] == 11:
                if i == 0:
                    listy.append("撞命關")
                elif i == 1:
                    listm.append("撞命關")
                elif i == 2:
                    listd.append("撞命關")
                else:
                    listh.append("撞命關")
    shansha[3] = shansha[3] + listh
    listh.clear()
    listd.clear()
    listy.clear()
    listm.clear()

    for i in range(4):
        if z[i] == 6:
            if z[1] == 2 or z[1] == 3:
                if i == 0:
                    listy.append("直難關")
                elif i == 1:
                    listm.append("直難關")
                elif i == 2:
                    listd.append("直難關")
                else:
                    listh.append("直難關")
        if z[i] == 7:
            if z[1] == 4 or z[1] == 5:
                if i == 0:
                    listy.append("直難關")
                elif i == 1:
                    listm.append("直難關")
                elif i == 2:
                    listd.append("直難關")
                else:
                    listh.append("直難關")
        if z[i] == 10 or z[i] == 3:
            if z[1] == 6 or z[1] == 7:
                if i == 0:
                    listy.append("直難關")
                elif i == 1:
                    listm.append("直難關")
                elif i == 2:
                    listd.append("直難關")
                else:
                    listh.append("直難關")
        if z[i] == 8 or z[i] == 5:
            if z[1] == 8 or z[1] == 9:
                if i == 0:
                    listy.append("直難關")
                elif i == 1:
                    listm.append("直難關")
                elif i == 2:
                    listd.append("直難關")
                else:
                    listh.append("直難關")
        if z[i] == 3 or z[i] == 2:
            if z[1] == 10 or z[1] == 11:
                if i == 0:
                    listy.append("直難關")
                elif i == 1:
                    listm.append("直難關")
                elif i == 2:
                    listd.append("直難關")
                else:
                    listh.append("直難關")
        if z[i] == 9 or z[i] == 4:
            if z[1] == 0 or z[1] == 1:
                if i == 0:
                    listy.append("直難關")
                elif i == 1:
                    listm.append("直難關")
                elif i == 2:
                    listd.append("直難關")
                else:
                    listh.append("直難關")

    for i in range(4):
        if z[i] == 10 or z[i] == 7:
            if z[1] == 2 or z[1] == 3 or z[1] == 4:
                if i == 0:
                    listy.append("水火關")
                elif i == 1:
                    listm.append("水火關")
                elif i == 2:
                    listd.append("水火關")
                else:
                    listh.append("水火關")
        if z[i] == 4 or z[i] == 1:
            if z[1] == 6 or z[1] == 5 or z[1] == 7 or z[1] == 8 or z[1] == 9 or z[1] == 10:
                if i == 0:
                    listy.append("水火關")
                elif i == 1:
                    listm.append("水火關")
                elif i == 2:
                    listd.append("水火關")
                else:
                    listh.append("水火關")
        if z[i] == 4 or z[i] == 7:
            if z[1] == 11 or z[1] == 0 or z[1] == 1:
                if i == 0:
                    listy.append("水火關")
                elif i == 1:
                    listm.append("水火關")
                elif i == 2:
                    listd.append("水火關")
                else:
                    listh.append("水火關")

    for i in range(4):
        if z[i] == 8 or z[i] == 2:
            if z[1] == 2 or z[1] == 3 or z[1] == 4:
                if i == 0:
                    listy.append("深水關")
                elif i == 1:
                    listm.append("深水關")
                elif i == 2:
                    listd.append("深水關")
                else:
                    listh.append("深水關")
        if z[i] == 7:
            if z[1] == 5 or z[1] == 6 or z[1] == 7:
                if i == 0:
                    listy.append("深水關")
                elif i == 1:
                    listm.append("深水關")
                elif i == 2:
                    listd.append("深水關")
                else:
                    listh.append("深水關")
        if z[i] == 9:
            if z[1] == 8 or z[1] == 9 or z[1] == 10:
                if i == 0:
                    listy.append("深水關")
                elif i == 1:
                    listm.append("深水關")
                elif i == 2:
                    listd.append("深水關")
                else:
                    listh.append("深水關")
        if z[i] == 1:
            if z[1] == 11 or z[1] == 0 or z[1] == 1:
                if i == 0:
                    listy.append("深水關")
                elif i == 1:
                    listm.append("深水關")
                elif i == 2:
                    listd.append("深水關")
                else:
                    listh.append("深水關")

    for i in range(4):
        if z[i] == 11 or z[i] == 5:
            if z[1] == 2 or z[1] == 8:
                if i == 0:
                    listy.append("四柱關")
                elif i == 1:
                    listm.append("四柱關")
                elif i == 2:
                    listd.append("四柱關")
                else:
                    listh.append("四柱關")
        if z[i] == 10 or z[i] == 4:
            if z[1] == 3 or z[1] == 9:
                if i == 0:
                    listy.append("四柱關")
                elif i == 1:
                    listm.append("四柱關")
                elif i == 2:
                    listd.append("四柱關")
                else:
                    listh.append("四柱關")
        if z[i] == 9 or z[i] == 3:
            if z[1] == 4 or z[1] == 10:
                if i == 0:
                    listy.append("四柱關")
                elif i == 1:
                    listm.append("四柱關")
                elif i == 2:
                    listd.append("四柱關")
                else:
                    listh.append("四柱關")
        if z[i] == 8 or z[i] == 2:
            if z[1] == 5 or z[1] == 11:
                if i == 0:
                    listy.append("四柱關")
                elif i == 1:
                    listm.append("四柱關")
                elif i == 2:
                    listd.append("四柱關")
                else:
                    listh.append("四柱關")
        if z[i] == 7 or z[i] == 1:
            if z[1] == 6 or z[1] == 0:
                if i == 0:
                    listy.append("四柱關")
                elif i == 1:
                    listm.append("四柱關")
                elif i == 2:
                    listd.append("四柱關")
                else:
                    listh.append("四柱關")
        if z[i] == 6 or z[i] == 0:
            if z[1] == 7 or z[1] == 1:
                if i == 0:
                    listy.append("四柱關")
                elif i == 1:
                    listm.append("四柱關")
                elif i == 2:
                    listd.append("四柱關")
                else:
                    listh.append("四柱關")

    for i in range(4):
        if z[i] == 4 or z[i] == 10 or z[i] == 9:
            if z[1] == 2 or z[1] == 3 or z[1] == 4:
                if i == 0:
                    listy.append("將軍箭")
                elif i == 1:
                    listm.append("將軍箭")
                elif i == 2:
                    listd.append("將軍箭")
                else:
                    listh.append("將軍箭")
        if z[i] == 0 or z[i] == 3 or z[i] == 7:
            if z[1] == 5 or z[1] == 6 or z[1] == 7:
                if i == 0:
                    listy.append("將軍箭")
                elif i == 1:
                    listm.append("將軍箭")
                elif i == 2:
                    listd.append("將軍箭")
                else:
                    listh.append("將軍箭")
        if z[i] == 1 or z[i] == 2 or z[i] == 6:
            if z[1] == 8 or z[1] == 9 or z[1] == 10:
                if i == 0:
                    listy.append("將軍箭")
                elif i == 1:
                    listm.append("將軍箭")
                elif i == 2:
                    listd.append("將軍箭")
                else:
                    listh.append("將軍箭")
        if z[i] == 5 or z[i] == 8 or z[i] == 11:
            if z[1] == 11 or z[1] == 0 or z[1] == 1:
                if i == 0:
                    listy.append("將軍箭")
                elif i == 1:
                    listm.append("將軍箭")
                elif i == 2:
                    listd.append("將軍箭")
                else:
                    listh.append("將軍箭")

    for i in range(4):
        if z[i] == 4:
            if z[1] == 2 or z[1] == 3 or z[1] == 4:
                if i == 0:
                    listy.append("浴盆關")
                elif i == 1:
                    listm.append("浴盆關")
                elif i == 2:
                    listd.append("浴盆關")
                else:
                    listh.append("浴盆關")
        if z[i] == 7:
            if z[1] == 5 or z[1] == 6 or z[1] == 7:
                if i == 0:
                    listy.append("浴盆關")
                elif i == 1:
                    listm.append("浴盆關")
                elif i == 2:
                    listd.append("浴盆關")
                else:
                    listh.append("浴盆關")
        if z[i] == 10:
            if z[1] == 8 or z[1] == 9 or z[1] == 10:
                if i == 0:
                    listy.append("浴盆關")
                elif i == 1:
                    listm.append("浴盆關")
                elif i == 2:
                    listd.append("浴盆關")
                else:
                    listh.append("浴盆關")
        if z[i] == 1:
            if z[1] == 11 or z[1] == 0 or z[1] == 1:
                if i == 0:
                    listy.append("浴盆關")
                elif i == 1:
                    listm.append("浴盆關")
                elif i == 2:
                    listd.append("浴盆關")
                else:
                    listh.append("浴盆關")

    for i in range(4):
        if z[i] == 2:
            if z[1] == 2:
                if i == 0:
                    listy.append("斷腸關")
                elif i == 1:
                    listm.append("斷腸關")
                elif i == 2:
                    listd.append("斷腸關")
                else:
                    listh.append("斷腸關")
        if z[i] == 3:
            if z[1] == 3:
                if i == 0:
                    listy.append("斷腸關")
                elif i == 1:
                    listm.append("斷腸關")
                elif i == 2:
                    listd.append("斷腸關")
                else:
                    listh.append("斷腸關")
        if z[i] == 8:
            if z[1] == 4:
                if i == 0:
                    listy.append("斷腸關")
                elif i == 1:
                    listm.append("斷腸關")
                elif i == 2:
                    listd.append("斷腸關")
                else:
                    listh.append("斷腸關")
        if z[i] == 1:
            if z[1] == 5:
                if i == 0:
                    listy.append("斷腸關")
                elif i == 1:
                    listm.append("斷腸關")
                elif i == 2:
                    listd.append("斷腸關")
                else:
                    listh.append("斷腸關")
        if z[i] == 10:
            if z[1] == 6:
                if i == 0:
                    listy.append("斷腸關")
                elif i == 1:
                    listm.append("斷腸關")
                elif i == 2:
                    listd.append("斷腸關")
                else:
                    listh.append("斷腸關")
        if z[i] == 9:
            if z[1] == 7:
                if i == 0:
                    listy.append("斷腸關")
                elif i == 1:
                    listm.append("斷腸關")
                elif i == 2:
                    listd.append("斷腸關")
                else:
                    listh.append("斷腸關")
        if z[i] == 4:
            if z[1] == 8:
                if i == 0:
                    listy.append("斷腸關")
                elif i == 1:
                    listm.append("斷腸關")
                elif i == 2:
                    listd.append("斷腸關")
                else:
                    listh.append("斷腸關")
        if z[i] == 5:
            if z[1] == 9:
                if i == 0:
                    listy.append("斷腸關")
                elif i == 1:
                    listm.append("斷腸關")
                elif i == 2:
                    listd.append("斷腸關")
                else:
                    listh.append("斷腸關")
        if z[i] == 6:
            if z[1] == 10:
                if i == 0:
                    listy.append("斷腸關")
                elif i == 1:
                    listm.append("斷腸關")
                elif i == 2:
                    listd.append("斷腸關")
                else:
                    listh.append("斷腸關")
        if z[i] == 7:
            if z[1] == 11:
                if i == 0:
                    listy.append("斷腸關")
                elif i == 1:
                    listm.append("斷腸關")
                elif i == 2:
                    listd.append("斷腸關")
                else:
                    listh.append("斷腸關")
        if z[i] == 11:
            if z[1] == 0:
                if i == 0:
                    listy.append("斷腸關")
                elif i == 1:
                    listm.append("斷腸關")
                elif i == 2:
                    listd.append("斷腸關")
                else:
                    listh.append("斷腸關")
        if z[i] == 0:
            if z[1] == 1:
                if i == 0:
                    listy.append("斷腸關")
                elif i == 1:
                    listm.append("斷腸關")
                elif i == 2:
                    listd.append("斷腸關")
                else:
                    listh.append("斷腸關")

    for i in range(4):
        if z[i] == 7 or z[i] == 1:
            if z[1] == 2 or z[1] == 3 or z[1] == 4:
                if i == 0:
                    listy.append("閻王關")
                elif i == 1:
                    listm.append("閻王關")
                elif i == 2:
                    listd.append("閻王關")
                else:
                    listh.append("閻王關")
        if z[i] == 10 or z[i] == 4:
            if z[1] == 5 or z[1] == 6 or z[1] == 7:
                if i == 0:
                    listy.append("閻王關")
                elif i == 1:
                    listm.append("閻王關")
                elif i == 2:
                    listd.append("閻王關")
                else:
                    listh.append("閻王關")
        if z[i] == 6 or z[i] == 0:
            if z[1] == 8 or z[1] == 9 or z[1] == 10:
                if i == 0:
                    listy.append("閻王關")
                elif i == 1:
                    listm.append("閻王關")
                elif i == 2:
                    listd.append("閻王關")
                else:
                    listh.append("閻王關")
        if z[i] == 3 or z[i] == 2:
            if z[1] == 11 or z[1] == 0 or z[1] == 1:
                if i == 0:
                    listy.append("閻王關")
                elif i == 1:
                    listm.append("閻王關")
                elif i == 2:
                    listd.append("閻王關")
                else:
                    listh.append("閻王關")

    for i in range(4):
        if z[i] == 0 or z[i] == 2 or z[i] == 9:
            if z[1] == 2 or z[1] == 3 or z[1] == 4:
                if i == 0:
                    listy.append("無情關")
                elif i == 1:
                    listm.append("無情關")
                elif i == 2:
                    listd.append("無情關")
                else:
                    listh.append("無情關")
        if z[i] == 5 or z[i] == 11 or z[i] == 10:
            if z[1] == 5 or z[1] == 6 or z[1] == 7:
                if i == 0:
                    listy.append("無情關")
                elif i == 1:
                    listm.append("無情關")
                elif i == 2:
                    listd.append("無情關")
                else:
                    listh.append("無情關")
        if z[i] == 1 or z[i] == 8:
            if z[1] == 8 or z[1] == 9 or z[1] == 10:
                if i == 0:
                    listy.append("無情關")
                elif i == 1:
                    listm.append("無情關")
                elif i == 2:
                    listd.append("無情關")
                else:
                    listh.append("無情關")
        if z[i] == 6 or z[i] == 0:
            if z[1] == 11 or z[1] == 0 or z[1] == 1:
                if i == 0:
                    listy.append("無情關")
                elif i == 1:
                    listm.append("無情關")
                elif i == 2:
                    listd.append("無情關")
                else:
                    listh.append("無情關")

    for i in range(4):
        if z[i] == 6 or z[i] == 1 or z[i] == 10 or z[i] == 4:
            if z[2] == 2 or z[2] == 5 or z[2] == 8 or z[2] == 11:
                if i == 0:
                    listy.append("百日關")
                elif i == 1:
                    listm.append("百日關")
                elif i == 2:
                    listd.append("百日關")
                else:
                    listh.append("百日關")
        if z[i] == 9 or z[i] == 3 or z[i] == 6 or z[i] == 0:
            if z[2] == 4 or z[2] == 7 or z[2] == 10 or z[2] == 1:
                if i == 0:
                    listy.append("百日關")
                elif i == 1:
                    listm.append("百日關")
                elif i == 2:
                    listd.append("百日關")
                else:
                    listh.append("百日關")
        if z[i] == 11 or z[i] == 5 or z[i] == 8 or z[i] == 2:
            if z[2] == 3 or z[2] == 6 or z[2] == 9 or z[2] == 0:
                if i == 0:
                    listy.append("百日關")
                elif i == 1:
                    listm.append("百日關")
                elif i == 2:
                    listd.append("百日關")
                else:
                    listh.append("百日關")

    for i in range(4):
        if z[i] == 5 or z[i] == 1:
            if z[1] == 2 or z[1] == 3 or z[1] == 4:
                if i == 0:
                    listy.append("四季關")
                elif i == 1:
                    listm.append("四季關")
                elif i == 2:
                    listd.append("四季關")
                else:
                    listh.append("四季關")
        if z[i] == 8 or z[i] == 4:
            if z[1] == 5 or z[1] == 6 or z[1] == 7:
                if i == 0:
                    listy.append("四季關")
                elif i == 1:
                    listm.append("四季關")
                elif i == 2:
                    listd.append("四季關")
                else:
                    listh.append("四季關")
        if z[i] == 11 or z[i] == 7:
            if z[1] == 8 or z[1] == 9 or z[1] == 10:
                if i == 0:
                    listy.append("四季關")
                elif i == 1:
                    listm.append("四季關")
                elif i == 2:
                    listd.append("四季關")
                else:
                    listh.append("四季關")
        if z[i] == 10 or z[i] == 2:
            if z[1] == 11 or z[1] == 0 or z[1] == 1:
                if i == 0:
                    listy.append("四季關")
                elif i == 1:
                    listm.append("四季關")
                elif i == 2:
                    listd.append("四季關")
                else:
                    listh.append("四季關")

    for i in range(4):
        if z[i] == 8:
            if z[1] == 2 or z[1] == 8:
                if i == 0:
                    listy.append("金鎖關")
                elif i == 1:
                    listm.append("金鎖關")
                elif i == 2:
                    listd.append("金鎖關")
                else:
                    listh.append("金鎖關")
        if z[i] == 9:
            if z[1] == 3 or z[1] == 9:
                if i == 0:
                    listy.append("金鎖關")
                elif i == 1:
                    listm.append("金鎖關")
                elif i == 2:
                    listd.append("金鎖關")
                else:
                    listh.append("金鎖關")
        if z[i] == 10:
            if z[1] == 4 or z[1] == 10:
                if i == 0:
                    listy.append("金鎖關")
                elif i == 1:
                    listm.append("金鎖關")
                elif i == 2:
                    listd.append("金鎖關")
                else:
                    listh.append("金鎖關")
        if z[i] == 11:
            if z[1] == 5 or z[1] == 11:
                if i == 0:
                    listy.append("金鎖關")
                elif i == 1:
                    listm.append("金鎖關")
                elif i == 2:
                    listd.append("金鎖關")
                else:
                    listh.append("金鎖關")
        if z[i] == 0:
            if z[1] == 6 or z[1] == 0:
                if i == 0:
                    listy.append("金鎖關")
                elif i == 1:
                    listm.append("金鎖關")
                elif i == 2:
                    listd.append("金鎖關")
                else:
                    listh.append("金鎖關")
        if z[i] == 1:
            if z[1] == 7 or z[1] == 1:
                if i == 0:
                    listy.append("金鎖關")
                elif i == 1:
                    listm.append("金鎖關")
                elif i == 2:
                    listd.append("金鎖關")
                else:
                    listh.append("金鎖關")
    shansha[0] = shansha[0] + listy
    shansha[2] = shansha[2] + listd
    shansha[3] = shansha[3] + listh


def place(table, x, y, str, isVerticle, size=15, setWid=False):
    if isVerticle:
        str = '\n'.join(str[i] for i in range(0, len(str)))
    table.cell(x, y).text = str
    table.cell(
        x, y).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    for run in table.cell(x, y).paragraphs[0].runs:
        font = run.font
        font.size = Pt(size)
    if setWid:
        table.cell(x, y).width = Inches(2)


def makeWord():
    doc = Document()  # 新建檔案
    table = doc.add_table(rows=9, cols=5, style='Table Grid')
    row = 0
    table.cell(row, 0).merge(table.cell(row, 4))
    table.cell(1, 0).merge(table.cell(1, 4))
    place(table, row, 0, f"{name}的八字命盤", False)
    row += 1
    if isLunar:
        place(table, row, 0,
              f"農曆 {year-1911} 年 {mon} 月 {day} 日 {hr} 時 生", False)
    else:
        place(table, row, 0,
              f"民國 {year-1911} 年 {mon} 月 {day} 日 {hr} 時 生", False)
    row += 1
    place(table, row, 4, "主星", True, 12)
    for i in range(3, -1, -1):
        place(table, row, i, mainStar[3-i], True)
    row += 1
    place(table, row, 4, "八字", True, 12)
    for i in range(3, -1, -1):
        toColor = {"火": RGBColor(255, 55, 55),
                   "金": RGBColor(255, 227, 132),
                   "木": RGBColor(34, 139, 34),
                   "水": RGBColor(30, 144, 255),
                   "土": RGBColor(210, 180, 140),
                   }
        table.cell(row, i).text = eightword_f[3-i]
        table.cell(row, i).add_paragraph('    '+eightword_f_elem[3-i])
        table.cell(row, i).add_paragraph(eightword_s[3-i])
        table.cell(row, i).add_paragraph('    '+eightword_s_elem[3-i])
        for j in table.cell(row, i).paragraphs:
            j.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.cell(row, i).paragraphs[0].runs[0].font.size = Pt(20)
        table.cell(row, i).paragraphs[0].runs[0].font.bold = True
        table.cell(row, i).paragraphs[1].runs[0].font.size = Pt(10)
        table.cell(
            row, i).paragraphs[1].runs[0].font.color.rgb = toColor[eightword_f_elem[3-i]]
        table.cell(row, i).paragraphs[2].runs[0].font.size = Pt(20)
        table.cell(row, i).paragraphs[2].runs[0].font.bold = True
        table.cell(row, i).paragraphs[3].runs[0].font.size = Pt(10)
        table.cell(
            row, i).paragraphs[3].runs[0].font.color.rgb = toColor[eightword_s_elem[3-i]]
    row += 1
    place(table, row, 4, "藏", True, 12)
    for i in range(3, -1, -1):
        tmp = ""
        for j in range(3):
            if bury[3-i][j] != "":
                tmp += bury[3-i][j]
            else:
                tmp += '　'
            tmp += ' '
        place(table, row, i, tmp, False)
    row += 1
    place(table, row, 4, "副星", True, 12)
    for i in range(3, -1, -1):
        tmp = ""
        for j in range(2):
            for k in range(3):
                if secondStar[3-i][k] != '':
                    tmp += secondStar[3-i][k][j]
                else:
                    tmp += '　'
                tmp += ' '
            tmp += '\n'
        place(table, row, i, tmp, False)
    row += 1
    place(table, row, 4, "運", True, 12)
    for i in range(3, -1, -1):
        place(table, row, i, fate[3-i], True)
    row += 1
    place(table, row, 4, "神煞", True, 12)
    for i in range(3, -1, -1):
        tmp = ""
        for j in shansha[3-i]:
            tmp += j+'\n'
        place(table, row, i, tmp, False, 12)
    row += 1
    if len(othershansha) != 0:
        table.add_row()
        table.cell(row, 0).merge(table.cell(row, 4))
        tmp = ""
        for j in othershansha:
            tmp += j+'  '
        place(table, row, 0, f"其他神煞: {tmp}", False)
        row += 1
    table.cell(row, 0).merge(table.cell(row, 1))
    place(table, row, 0, loveGod, False)
    table.cell(row, 2).merge(table.cell(row, 3))
    place(table, row, 2, body, False)

    table2 = doc.add_table(rows=3, cols=11, style='Table Grid')
    table2.cell(0, 10).merge(table2.cell(1, 10))
    place(table2, 0, 10, "大運", False, 12)
    bf = int(bigFate)
    for i in range(9, -1, -1):
        place(table2, 0, i, str(bf), False)
        bf += 10
    for i in range(9, -1, -1):
        place(table2, 1, i, bigFateYears[9-i], True)
    table2.cell(2, 0).merge(table2.cell(2, 10))
    place(table2, 2, 0, bigFateText, False)
    doc.add_page_break()
    table3 = doc.add_table(rows=37, cols=10, style='Table Grid')
    table3.cell(0, 0).merge(table3.cell(0, 9))
    place(table3, 0, 0, "流年", False, 10)
    row = 1
    cnt = int(bigFate)
    for i in range(1, 37, 4):
        for j in range(10):
            place(table3, i, j, str(cnt+year - 1912), False, 10)
            place(table3, i+1, j, str(cnt), False, 10)
            place(table3, i+2, j, flowYears[cnt], True, 10)
            cnt += 1
    doc.add_page_break()
    table5 = doc.add_table(rows=1, cols=2, style='Table Grid')
    row = 0
    place(table5, row, 0, "神煞", False, 15, True)
    place(table5, row, 1, "說明", False, 15, True)
    row += 1
    for k, v in shanshaExplain.items():
        table5.add_row()
        place(table5, row, 0, k, False, 12, True)
        place(table5, row, 1, v, False, 12, True)
        row += 1
    doc.add_page_break()
    table4 = doc.add_table(rows=11, cols=2, style='Table Grid')
    row = 0

    place(table4, row, 0, "個性", False, 15, True)
    place(table4, row, 1, "說明", False, 15, True)
    row += 1
    for k, v in personality.items():
        place(table4, row, 0, k, False, 12, True)
        place(table4, row, 1, v, False, 12, True)
        row += 1
    table4.cell(row, 0).merge(table4.cell(row, 1))
    place(table4, row, 0, " ", False, 15, True)
    row += 1
    place(table4, row, 0, "馬花庫", False, 15, True)
    place(table4, row, 1, "說明", False, 15, True)
    row += 1
    for k, v in horseFlower.items():
        place(table4, row, 0, k, False, 12, True)
        place(table4, row, 1, v, False, 12, True)
        row += 1

    doc.save(f'.\\Words\\{name}.docx')


def send(name):
    # coding=utf-8
    win.destroy()
    file = '.\\Words\\'+str(name)+'.docx'
    server_address = 'smtp.gmail.com'
    port = 587
    send_user = 'why33551@gmail.com'
    receive_users = 'mama33551@gmail.com'
    subject = f'{name}的八字命盤'
    password = "tyevkckncblwgnut"

    mail_type = '1'

    msg = MIMEMultipart()
    msg['Subject'] = subject
    msg['From'] = send_user
    msg['To'] = receive_users

    part_text = MIMEText('請接收以下檔案')
    msg.attach(part_text)

    part_attach1 = MIMEApplication(open(file, 'rb').read())
    part_attach1.add_header('Content-Disposition',
                            'attachment', filename=file)
    msg.attach(part_attach1)

    smtp = smtplib.SMTP(server_address, port)
    smtp.ehlo()  # 驗證SMTP伺服器
    smtp.starttls()
    smtp.login(send_user, password)
    smtp.sendmail(send_user, receive_users, msg.as_string())
    print('郵件傳送成功！')
    quit()


def sendGUI():
    global name
    global visited
    global win
    win.destroy()
    win = tk.Tk()
    setupWin()
    name_l = tk.Label(text="傳送", font="微軟正黑體 20")
    name_c = ttk.Combobox(win, font="微軟正黑體 20", state="readonly")
    recnames = []
    for file in glob.glob(".\\Words\\*.docx"):
        recnames.append(str(file)[8:-5])
    name_c['value'] = recnames
    name_c.current(0)
    name_b = tk.Button(
        text="確定", command=lambda: send(name_c.get()), font="微軟正黑體 15")
    name_l.grid(row=0, column=0)
    name_c.grid(row=0, column=1)
    name_b.grid(row=0, column=2)


def dele(name):
    win.destroy()
    file = '.\\Words\\'+str(name)+'.docx'
    os.remove(file)
    quit()


def delGUI():
    global name
    global visited
    global win
    win.destroy()
    win = tk.Tk()
    setupWin()
    name_l = tk.Label(text="刪除", font="微軟正黑體 20")
    name_c = ttk.Combobox(win, font="微軟正黑體 20", state="readonly")
    recnames = []
    for file in glob.glob(".\\Words\\*.docx"):
        recnames.append(str(file)[8:-5])
    name_c['value'] = recnames
    name_c.current(0)
    name_b = tk.Button(
        text="確定", command=lambda: dele(name_c.get()), font="微軟正黑體 15")
    name_l.grid(row=0, column=0)
    name_c.grid(row=0, column=1)
    name_b.grid(row=0, column=2)


def menuGUI():
    setupWin()
    add_b = tk.Button(
        text="新增", command=GUI, font="微軟正黑體 15")
    add_b.pack()
    send_b = tk.Button(
        text="傳送E-mail", command=sendGUI, font="微軟正黑體 15")
    send_b.pack()
    del_b = tk.Button(
        text="刪除", command=delGUI, font="微軟正黑體 15", bg="red")
    del_b.pack()
    win.mainloop()


# main
name = " "
year = 0
mon = 0
day = 0
hr = 0
recData = []  # [year, name, isLunar, isMan, mon, day, hr, isLeapMonth]
recorded = False
isMan = True
isLunar = False
isLeapMonth = False
checkAll = False
win = tk.Tk()
menuGUI()

addNewData()
if not checkAll:
    exit()
mainStar = []
eightword_f = []
eightword_s = []
eightword_f_elem = []
eightword_s_elem = []
bury = [[], [], [], []]
secondStar = [[], [], [], []]
fate = []
shansha = [[], [], [], []]
othershansha = []
bigFate = ""
bigFateYears = []
bigFateText = ""
flowYears = []
flowYears.append(" ")  # index 1
loveGod = ""
body = ""
personality = {}
horseFlower = {}
shanshaExplain = {}

crawler()
adjust()

makeWord()
